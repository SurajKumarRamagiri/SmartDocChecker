"""
Text Extraction Layer

Converts uploaded files (PDF, DOCX) to clean raw text.
"""

import io
import logging
from pathlib import Path
from typing import Union

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

import re as _re
from collections import Counter

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


def extract_text_from_pdf(file_content: Union[bytes, io.BytesIO]) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_content: PDF file content as bytes or BytesIO object
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If PDF extraction fails
    """
    if PdfReader is None:
        raise TextExtractionError("PyPDF2 is not installed. Install with: pip install PyPDF2")
    
    try:
        # Convert bytes to BytesIO if needed
        if isinstance(file_content, bytes):
            file_content = io.BytesIO(file_content)
        
        # Read PDF
        pdf_reader = PdfReader(file_content)
        
        # Extract text from all pages
        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                continue
        
        raw_text = "\n\n".join(text_parts)
        
        if not raw_text.strip():
            raise TextExtractionError("No text could be extracted from PDF")
        
        # ── PDF-specific cleanup ──
        raw_text = _clean_pdf_text(raw_text, len(pdf_reader.pages))
        
        return raw_text.strip()
        
    except Exception as e:
        if isinstance(e, TextExtractionError):
            raise
        raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")


def _clean_pdf_text(raw_text: str, num_pages: int) -> str:
    """
    PDF-specific cleanup:
    - Remove repeated header / footer lines that appear on most pages
      (handles page-number variations like "Company Name | 3 Report Title")
    - Strip standalone page numbers and running titles
    """
    lines = raw_text.split('\n')
    if num_pages < 2 or len(lines) < 10:
        return raw_text

    # ── 1. Normalise page-number tokens before counting ──
    # Replace isolated numbers (possibly page numbers embedded in running
    # headers) with a placeholder so "Title | 3 Foo" and "Title | 5 Foo"
    # are treated as the same line.

    def _normalise_line(line: str) -> str:
        """Collapse page-number-like tokens so header variants match."""
        s = line.strip()
        # "| 3 " or "| 12 " (pipe + number seen in academic running headers)
        s = _re.sub(r'\|\s*\d{1,4}\s', '| # ', s)
        # Standalone leading/trailing numbers: "3 Company…" or "…Title 12"
        s = _re.sub(r'(?:^|\s)\d{1,4}(?:\s|$)', ' # ', s)
        return ' '.join(s.split())

    norm_counts = Counter()
    for line in lines:
        stripped = line.strip()
        if stripped:
            norm_counts[_normalise_line(stripped)] += 1

    # Lines whose normalised form appears on >40 % of pages are headers/footers
    repeat_threshold = max(2, int(num_pages * 0.4))
    repeated_norms = {norm for norm, cnt in norm_counts.items()
                      if cnt >= repeat_threshold and len(norm) < 140}

    cleaned = []
    for line in lines:
        stripped = line.strip()
        if stripped and _normalise_line(stripped) in repeated_norms:
            continue
        # Strip standalone page numbers ("1", "Page 5", "- 3 -")
        if _re.match(r'^\s*[-–—]?\s*(?:page\s+)?\d{1,4}\s*[-–—]?\s*$', line, _re.IGNORECASE):
            continue
        cleaned.append(line)

    return '\n'.join(cleaned)


def extract_text_from_docx(file_content: Union[bytes, io.BytesIO]) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_content: DOCX file content as bytes or BytesIO object
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If DOCX extraction fails
    """
    if DocxDocument is None:
        raise TextExtractionError("python-docx is not installed. Install with: pip install python-docx")
    
    try:
        # Convert bytes to BytesIO if needed
        if isinstance(file_content, bytes):
            file_content = io.BytesIO(file_content)
        
        # Read DOCX
        doc = DocxDocument(file_content)
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables – but skip numeric/data tables
        for table in doc.tables:
            # Sample all cells to decide if the table is worth extracting
            total_cells = 0
            numeric_cells = 0
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        total_cells += 1
                        # Count cells that are purely numeric/currency/date
                        if _re.match(r'^[\d\s\.\,\$\%\€\£\(\)\-\+\/\|:]+$', txt):
                            numeric_cells += 1
            # Skip table if >50% of cells are numeric data
            if total_cells > 0 and (numeric_cells / total_cells) > 0.5:
                continue
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        raw_text = "\n\n".join(text_parts)
        
        if not raw_text.strip():
            raise TextExtractionError("No text could be extracted from DOCX")
        
        return raw_text.strip()
        
    except Exception as e:
        if isinstance(e, TextExtractionError):
            raise
        raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt(file_content: Union[bytes, io.BytesIO]) -> str:
    """
    Extract text from a TXT file.
    
    Args:
        file_content: TXT file content as bytes or BytesIO object
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If TXT extraction fails
    """
    try:
        # Convert bytes to BytesIO if needed
        if isinstance(file_content, bytes):
            file_content = io.BytesIO(file_content)
        
        # Read text file
        raw_text = file_content.read().decode('utf-8')
        
        if not raw_text.strip():
            raise TextExtractionError("No text could be extracted from TXT file")
        
        return raw_text.strip()
        
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            if isinstance(file_content, bytes):
                file_content = io.BytesIO(file_content)
            else:
                file_content.seek(0)
            
            raw_text = file_content.read().decode('latin-1')
            
            if not raw_text.strip():
                raise TextExtractionError("No text could be extracted from TXT file")
            
            return raw_text.strip()
        except Exception as e:
            raise TextExtractionError(f"Failed to decode TXT file: {str(e)}")
    except Exception as e:
        if isinstance(e, TextExtractionError):
            raise
        raise TextExtractionError(f"Failed to extract text from TXT: {str(e)}")


def extract_text(file_content: Union[bytes, io.BytesIO], filename: str) -> str:
    """
    Extract text from a file based on its extension.
    
    Args:
        file_content: File content as bytes or BytesIO object
        filename: Name of the file (used to determine file type)
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If extraction fails or file type is unsupported
    """
    # Get file extension
    ext = Path(filename).suffix.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_content)
    elif ext == '.docx':
        return extract_text_from_docx(file_content)
    elif ext == '.doc':
        raise TextExtractionError(
            "Legacy .doc format is not supported. "
            "Please convert to .docx or PDF and re-upload."
        )
    elif ext == '.txt':
        return extract_text_from_txt(file_content)
    else:
        raise TextExtractionError(f"Unsupported file type: {ext}. Supported types: .pdf, .docx, .txt")


def clean_text(raw_text: str) -> str:
    """
    Clean extracted text by removing excessive whitespace and normalizing line breaks.
    
    Args:
        raw_text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    lines = [line.strip() for line in raw_text.split('\n')]
    
    # Remove empty lines
    lines = [line for line in lines if line]
    
    # Join with single line breaks
    cleaned = '\n'.join(lines)
    
    # Normalize multiple spaces to single space
    cleaned = _re.sub(r' +', ' ', cleaned)
    
    return cleaned


# Convenience function that extracts and cleans in one step
def extract_and_clean_text(file_content: Union[bytes, io.BytesIO], filename: str) -> str:
    """
    Extract and clean text from a file.
    
    Args:
        file_content: File content as bytes or BytesIO object
        filename: Name of the file (used to determine file type)
        
    Returns:
        Cleaned extracted text as a string
        
    Raises:
        TextExtractionError: If extraction fails or file type is unsupported
    """
    raw_text = extract_text(file_content, filename)
    return clean_text(raw_text)


# ── FastAPI UploadFile compatibility ──
async def extract_text_from_upload(file) -> str:
    """
    Extract text from a FastAPI UploadFile object.
    
    This function provides backward compatibility with existing code
    that uses FastAPI's UploadFile type.
    
    Args:
        file: FastAPI UploadFile object
        
    Returns:
        Extracted and cleaned text as a string
        
    Raises:
        TextExtractionError: If extraction fails
    """
    try:
        # Read file content
        content = await file.read()
        
        # Reset file pointer for potential reuse
        await file.seek(0)
        
        # Extract and clean text
        return extract_and_clean_text(content, file.filename)
        
    except Exception as e:
        if isinstance(e, TextExtractionError):
            raise
        raise TextExtractionError(f"Failed to extract text from upload: {str(e)}")

